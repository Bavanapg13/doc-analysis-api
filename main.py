"""
AI-Powered Document Analysis & Extraction API
Supports: PDF, DOCX, Images (OCR)
"""

import os
import io
import re
import json
import tempfile
import logging
from typing import Optional

import anthropic
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document
from fastapi import FastAPI, File, UploadFile, Header, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="AI Document Analysis API",
    description="Extract, analyse, and summarise PDF, DOCX, and image documents",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.getenv("DOC_API_KEY", "doc-api-key-2024-secure-xyz")

def verify_api_key(authorization: Optional[str] = Header(None), x_api_key: Optional[str] = Header(None)):
    key = None
    if authorization:
        key = authorization.replace("Bearer ", "").strip()
    elif x_api_key:
        key = x_api_key.strip()
    if key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing API key")
    return key

anthropic_client = anthropic.Anthropic()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text(layout=True) or ""
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")
            else:
                img = page.to_image(resolution=200).original
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    text_parts.append(f"[Page {page_num} (OCR)]\n{ocr_text}")
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            if para.style.name.startswith("Heading"):
                parts.append(f"\n## {para.text.strip()}")
            else:
                parts.append(para.text.strip())
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
    return "\n".join(parts)


def extract_text_from_image(file_bytes: bytes) -> str:
    img = Image.open(io.BytesIO(file_bytes))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    text = pytesseract.image_to_string(img, config="--psm 6")
    return text.strip()


def detect_file_type(filename: str, content_type: str) -> str:
    name_lower = filename.lower()
    ct_lower = (content_type or "").lower()
    if name_lower.endswith(".pdf") or "pdf" in ct_lower:
        return "pdf"
    if name_lower.endswith(".docx") or "wordprocessingml" in ct_lower:
        return "docx"
    if any(name_lower.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif", ".webp"]):
        return "image"
    if "image" in ct_lower:
        return "image"
    return "unknown"


ANALYSIS_SYSTEM_PROMPT = """You are an expert document analysis AI. Given extracted document text, you will:
1. Write a concise, accurate summary (2-4 sentences)
2. Extract named entities categorised as: persons, organisations, locations, dates, monetary_amounts, other
3. Classify overall sentiment as: positive, negative, or neutral with a confidence score (0.0-1.0)

Respond ONLY with valid JSON matching this exact schema:
{
  "summary": "string",
  "entities": {
    "persons": ["string"],
    "organisations": ["string"],
    "locations": ["string"],
    "dates": ["string"],
    "monetary_amounts": ["string"],
    "other": ["string"]
  },
  "sentiment": {
    "label": "positive|negative|neutral",
    "score": 0.0
  }
}

Return ONLY the JSON object, no markdown fences, no extra text"""


def analyse_with_ai(text: str) -> dict:
    max_chars = 12000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[... document truncated ...]"

    response = anthropic_client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1500,
        system=ANALYSIS_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": f"Analyse this document text:\n\n{text}"}],
    )

    raw = response.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


@app.get("/")
async def root():
    return {
        "service": "AI Document Analysis API",
        "version": "1.0.0",
        "endpoints": {"POST /analyse": "Analyse a document", "GET /health": "Health check"},
        "supported_formats": ["pdf", "docx", "png", "jpg", "jpeg", "tiff", "bmp", "webp"],
    }


@app.get("/health")
async def health():
    return {"status": "ok", "service": "doc-analysis-api"}


@app.post("/analyse")
async def analyse_document(
    file: UploadFile = File(...),
    _key: str = Depends(verify_api_key),
):
    filename = file.filename or "document"
    content_type = file.content_type or ""
    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty file uploaded")

    file_type = detect_file_type(filename, content_type)
    if file_type == "unknown":
        raise HTTPException(status_code=415, detail=f"Unsupported file type: {filename}")

    try:
        if file_type == "pdf":
            extracted_text = extract_text_from_pdf(file_bytes)
        elif file_type == "docx":
            extracted_text = extract_text_from_docx(file_bytes)
        else:
            extracted_text = extract_text_from_image(file_bytes)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Text extraction failed: {str(exc)}")

    if not extracted_text.strip():
        raise HTTPException(status_code=422, detail="No text could be extracted from the document")

    try:
        analysis = analyse_with_ai(extracted_text)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"AI analysis failed: {str(exc)}")

    return {
        "filename": filename,
        "file_type": file_type,
        "word_count": len(extracted_text.split()),
        "summary": analysis.get("summary", ""),
        "entities": analysis.get("entities", {}),
        "sentiment": analysis.get("sentiment", {}),
    }


app.add_api_route("/analyze", analyse_document, methods=["POST"])

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=False)

